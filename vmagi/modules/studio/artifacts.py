"""
Fábrica de artefactos con bucle de observación (Plan MAGI 9.0 §5).

EL PATRÓN
=========
    ESPECIFICAR -> GENERAR -> EJECUTAR/RENDERIZAR -> OBSERVAR -> CRITICAR -> ITERAR

La clave es OBSERVAR. Un sistema que genera un juego y te lo entrega sin
haberlo arrancado ha *generado código de juego*; uno que lo arranca, captura un
fotograma y lo mira, ha *hecho un juego*. La diferencia no está en el modelo:
está en si hay un bucle que cierra sobre el resultado.

Es la misma idea que la verificación ejecutable del §2.5, aplicada a artefactos
que no son código: una imagen se mira, un documento se renderiza y se cuenta
las páginas, un juego se arranca en modo headless y se captura.
"""
from __future__ import annotations

import asyncio
import logging
import shutil
from dataclasses import dataclass, field
from enum import Enum
from pathlib import Path

from ...core.paths import python_executable

logger = logging.getLogger(__name__)

#: Lo que se dice cuando el .exe no tiene con qué ejecutar Python.
#: Ver `paths.python_executable`: dentro del bundle `sys.executable`
#: es el propio .exe y lanzarlo relanzaría MAGI.
_SIN_PYTHON = (
    "no hay un intérprete de Python en esta máquina. El .exe de MAGI no "
    "puede ejecutarlo por sí solo: dentro del bundle `sys.executable` es "
    "el propio .exe, así que lanzarlo relanzaría MAGI en vez de ejecutar "
    "esto. Instala Python y vuelve a intentarlo.")


class ArtifactKind(str, Enum):
    PROGRAM = "programa"
    GAME = "juego"
    IMAGE = "imagen"
    DOCUMENT = "documento"
    VIDEO = "video"
    DATA = "datos"


@dataclass
class Observation:
    """Lo que el sistema VE de su propio resultado."""
    ok: bool
    kind: ArtifactKind
    summary: str
    evidence: list[str] = field(default_factory=list)
    artifact_path: str | None = None
    screenshot: str | None = None
    problems: list[str] = field(default_factory=list)

    def render(self) -> str:
        head = f"[{'OK' if self.ok else 'FALLA'}] {self.kind.value}: {self.summary}"
        parts = [head]
        if self.evidence:
            parts += [f"  · {e}" for e in self.evidence]
        if self.problems:
            parts.append("  problemas observados:")
            parts += [f"  · {p}" for p in self.problems]
        if self.screenshot:
            parts.append(f"  captura: {self.screenshot}")
        return "\n".join(parts)

    def feedback(self) -> str:
        """Lo que vuelve al generador cuando algo no cuadra."""
        if self.ok and not self.problems:
            return ""
        return ("El artefacto se ha inspeccionado y presenta estos problemas. "
                "Corrígelos:\n" + "\n".join(f"- {p}" for p in self.problems))


# --------------------------------------------------------------- programas

#: nombres que se prueban como punto de entrada de un proyecto Python
ENTRY_CANDIDATES = ("main.py", "__main__.py", "app.py", "run.py", "start.py")

#: extensiones que se despachan a `observe_video` (§5.5). Antes ninguna estaba
#: reconocida y todas acababan en `observe_program`, es decir, ejecutándose.
VIDEO_EXTS = (".mp4", ".mkv", ".mov", ".webm", ".avi", ".gif", ".m4v")


async def observe_program(path: str | Path, *, entry: str = "",
                          timeout: int = 60) -> Observation:
    """Arranca el programa y mira si sobrevive."""
    _interprete = python_executable()
    p = Path(path)
    destino = p.name
    if not p.exists():
        return Observation(False, ArtifactKind.PROGRAM, "no existe",
                           problems=[f"{p} no existe"])

    if p.is_dir() and not entry:
        # Antes se construía `python <nombre-del-directorio>`, que fallaba con
        # rc=2 y un mensaje que no explicaba nada. Un directorio es un
        # proyecto: hay que buscar su punto de entrada.
        found = next((c for c in ENTRY_CANDIDATES if (p / c).exists()), None)
        if found is None:
            return Observation(
                False, ArtifactKind.PROGRAM, "sin punto de entrada",
                artifact_path=str(p),
                problems=[f"{p} es un directorio y no contiene ninguno de "
                          f"{', '.join(ENTRY_CANDIDATES)}. Indica `entry` con "
                          f"el comando de arranque."])
        destino = found

    # Solo hace falta un intérprete si NO nos han dado la orden de arranque:
    # un `entry` explícito puede ser `./juego` o `node app.js`. Cuando sí hace
    # falta y no lo hay, se dice — antes se componía la orden con el .exe de
    # MAGI y se relanzaba MAGI en vez de ejecutar el artefacto.
    if not entry and _interprete is None:
        return Observation(
            False, ArtifactKind.PROGRAM, "sin intérprete de Python",
            artifact_path=str(p), problems=[_SIN_PYTHON])

    cmd = entry or f'"{_interprete}" "{destino}"'
    cwd = p if p.is_dir() else p.parent
    rc, out = await _run(cmd, cwd, timeout)

    problems = []
    if rc != 0:
        problems.append(f"termina con código {rc}")
        tail = "\n".join(out.strip().splitlines()[-6:])
        if tail:
            problems.append(f"últimas líneas:\n{tail}")
    return Observation(rc == 0, ArtifactKind.PROGRAM,
                       f"ejecutado (rc={rc})",
                       evidence=[out[-800:]] if out else [],
                       artifact_path=str(p), problems=problems)


# ------------------------------------------------------------------ juegos

PYGAME_HARNESS = '''"""
Arnés de observación generado por MAGI (§5.2).

Arranca el juego en modo headless, avanza N fotogramas y guarda una captura.
Sin esto el sistema entrega "código de juego" sin saber si el jugador se
distingue del fondo, o si la pantalla sale en negro.
"""
import os, sys, importlib.util
os.environ.setdefault("SDL_VIDEODRIVER", "dummy")
os.environ.setdefault("SDL_AUDIODRIVER", "dummy")

import pygame

FRAMES = int(os.environ.get("MAGI_FRAMES", "120"))
SHOT = os.environ.get("MAGI_SHOT", "frame.png")
TARGET = os.environ.get("MAGI_TARGET", "main.py")

pygame.init()
_orig_flip, _orig_update = pygame.display.flip, pygame.display.update
state = {"frames": 0}


def _capture():
    surf = pygame.display.get_surface()
    if surf is not None:
        pygame.image.save(surf, SHOT)


def _cuenta():
    state["frames"] += 1
    if state["frames"] >= FRAMES:
        _capture()
        pygame.quit()
        raise SystemExit(0)


def _flip(*a, **kw):
    _cuenta()
    return _orig_flip(*a, **kw)


def _update(*a, **kw):
    # update() y flip() NO son intercambiables: flip() no acepta argumentos y
    # update(rect) sí. Aliasar los dos al mismo envoltorio hacía que cualquier
    # juego con dirty rects muriera con
    #     TypeError: pygame.display.flip() takes no arguments (1 given)
    # y el informe culpaba al juego de no dibujar nada.
    _cuenta()
    return _orig_update(*a, **kw)


pygame.display.flip = _flip
pygame.display.update = _update

spec = importlib.util.spec_from_file_location("__main__", TARGET)
mod = importlib.util.module_from_spec(spec)
try:
    spec.loader.exec_module(mod)
except SystemExit:
    pass
finally:
    if state["frames"] > 0:
        _capture()
    print(f"MAGI_FRAMES_RENDERED={state['frames']}")
'''


async def observe_game(project_dir: str | Path, *, entry: str = "main.py",
                       frames: int = 120, timeout: int = 90) -> Observation:
    """
    Arranca un juego Pygame en headless, avanza fotogramas y captura uno.

    Lo que esto permite: que Balthasar MIRE la captura con visión y diga "el
    jugador no se distingue del fondo" — una crítica imposible de hacer
    leyendo el código.
    """
    d = Path(project_dir)
    target = d / entry
    if not target.exists():
        return Observation(False, ArtifactKind.GAME, "sin punto de entrada",
                           problems=[f"no existe {target}"])

    try:
        import pygame  # noqa: F401
    except ImportError:
        return Observation(
            False, ArtifactKind.GAME, "pygame no instalado",
            problems=["pip install pygame para poder observar el juego. "
                      "Sin esto solo se puede revisar el código, no verlo."])

    harness = d / "_magi_harness.py"
    harness.write_text(PYGAME_HARNESS, encoding="utf-8")
    shot = d / "_magi_frame.png"
    # Borrar la captura anterior: si el juego revienta antes de dibujar, el
    # fichero viejo sigue ahí y el informe describe la ejecución EQUIVOCADA.
    shot.unlink(missing_ok=True)
    env = {"MAGI_FRAMES": str(frames), "MAGI_SHOT": str(shot),
           "MAGI_TARGET": str(target), "SDL_VIDEODRIVER": "dummy",
           "SDL_AUDIODRIVER": "dummy"}

    interprete = python_executable()
    if interprete is None:
        harness.unlink(missing_ok=True)
        return Observation(False, ArtifactKind.GAME, "sin intérprete de Python",
                           artifact_path=str(d), problems=[_SIN_PYTHON])
    try:
        rc, out = await _run(f'"{interprete}" "{harness.name}"', d,
                             timeout, env)
    finally:
        harness.unlink(missing_ok=True)

    rendered = 0
    for line in out.splitlines():
        if line.startswith("MAGI_FRAMES_RENDERED="):
            rendered = int(line.split("=")[1] or 0)

    problems = []
    if rendered == 0:
        problems.append("no se dibujó ni un fotograma: el juego no llega a "
                        "arrancar o nunca llama a display.flip()")
    if not shot.exists():
        problems.append("no se pudo capturar la pantalla")
    if rc != 0 and rendered == 0:
        problems.append(f"salida con código {rc}:\n" +
                        "\n".join(out.strip().splitlines()[-6:]))

    evidence = [f"código de salida {rc}"]
    if shot.exists():
        # El fallo que este bucle existe para cazar: el juego corre, dibuja
        # fotogramas y en pantalla no se ve nada porque todo es del mismo
        # color. Sin esto el informe decía OK y enterraba la pista en la
        # evidencia. Y sin Pillow ni siquiera se enteraba: ver `_mirar_imagen`.
        desc, malos = _mirar_imagen(
            shot,
            vacia="la pantalla es de un solo color: el juego dibuja pero no se "
                  "ve nada. Revisa que los elementos no sean del color del "
                  "fondo y que se dibujen dentro de los límites de la ventana.")
        evidence.append(desc)
        problems.extend(malos)

    return Observation(
        bool(rendered) and shot.exists() and not problems, ArtifactKind.GAME,
        f"{rendered} fotogramas dibujados",
        evidence=evidence, artifact_path=str(d),
        screenshot=str(shot) if shot.exists() else None,
        problems=problems)


# ------------------------------------------------------------------ imagen

def pillow_available() -> bool:
    """
    ¿Se puede MIRAR una imagen en esta máquina?

    Existe como función y no como `try: import PIL` disperso porque quien
    observa necesita distinguir dos cosas que se parecen mucho y significan lo
    contrario: «he mirado y está bien» y «no he podido mirar». Sin esta
    pregunta explícita, la segunda se colaba como la primera.
    """
    # `import PIL` NO basta: PIL es un paquete namespace que importa aunque su
    # extensión en C esté rota (rueda de arquitectura equivocada, desinstalación
    # a medias, bundle de PyInstaller sin `_imaging`). En ese estado
    # `pillow_available()` decía True, la guarda no saltaba y volvía el mismo
    # fallo un nivel más abajo. Se pregunta por la capacidad que se va a usar.
    try:
        from PIL import Image  # noqa: F401
    except Exception:
        return False
    return True


#: Lo que devuelve `_describe_image` cuando no hay con qué mirar. Quien observe
#: tiene que tratarlo como un PROBLEMA, no como una descripción.
SIN_PILLOW = "Pillow no instalado: no se puede inspeccionar la imagen"


def _mirar_imagen(path: str | Path, *, vacia: str) -> tuple[str, list[str]]:
    """
    Describe una imagen Y dice qué problemas hay, en una sola llamada.

    POR QUÉ EXISTE. `_describe_image` devuelve texto libre y cada observador
    tenía que ACORDARSE de buscar la subcadena correcta dentro. Tres de ellos
    solo buscaban "VACÍA", así que cuando faltaba Pillow —o la imagen era
    ilegible— la lista de problemas quedaba vacía y `ok` salía True: el
    observador certificaba como buena una imagen que nunca llegó a abrir.

    Buscar subcadenas en prosa es un acoplamiento que se rompe en silencio.
    Aquí se rompe una vez, en un sitio, y todos los llamadores heredan el
    arreglo. `vacia` es el mensaje propio de cada artefacto, que sí cambia.
    """
    desc = _describe_image(path)
    if desc == SIN_PILLOW:
        return desc, ["Pillow no instalado: la imagen NO se ha mirado. No se "
                      "sabe si está en negro o vacía. `pip install pillow` "
                      "para cerrar el bucle de observación."]
    if desc.startswith("imagen ilegible"):
        return desc, [desc]
    if "VACÍA" in desc:
        return desc, [vacia]
    return desc, []


def _describe_image(path: str | Path) -> str:
    """
    Descripción objetiva de una imagen SIN modelo de visión.

    Detecta el fallo más común de una captura de juego: la pantalla en negro,
    o una imagen de un solo color. Es barato y no gasta cuota.
    """
    if not pillow_available():
        return SIN_PILLOW
    try:
        from PIL import Image
    except ImportError:                           # pragma: no cover
        return SIN_PILLOW
    try:
        with Image.open(path) as im:
            im = im.convert("RGB")
            w, h = im.size
            colors = im.getcolors(maxcolors=1_000_000)
            if not colors:
                return f"{w}x{h}, muchos colores distintos"
            colors.sort(reverse=True)
            top_count, top_rgb = colors[0]
            share = top_count / (w * h)
            desc = (f"{w}x{h}, {len(colors)} colores; el dominante "
                    f"{top_rgb} ocupa el {share:.0%}")
            if share > 0.98:
                desc += "  <-- PANTALLA PRÁCTICAMENTE VACÍA"
            return desc
    except Exception as e:
        return f"imagen ilegible: {e}"


async def observe_image(path: str | Path) -> Observation:
    p = Path(path)
    if not p.exists():
        return Observation(False, ArtifactKind.IMAGE, "no existe",
                           problems=[f"{p} no existe"])
    # EL FALLO QUE ESTO CIERRA: sin Pillow, `_describe_image` devolvía «no se
    # puede inspeccionar», esa cadena no contenía ni "VACÍA" ni "ilegible", la
    # lista de problemas quedaba vacía y `ok` salía True. El observador
    # certificaba como buena una imagen que nunca llegó a abrir. Lo encontró la
    # simulación del entorno de CI, donde Pillow no estaba instalado.
    desc, problems = _mirar_imagen(
        p, vacia="la imagen es casi de un solo color: probablemente no se "
                 "dibujó nada")
    return Observation(not problems, ArtifactKind.IMAGE, desc,
                       artifact_path=str(p), screenshot=str(p),
                       problems=problems)


# -------------------------------------------------------------- documentos

#: extensiones de datos que se despachan a `observe_data`.
DATA_EXTS = (".csv", ".tsv", ".json", ".jsonl", ".ndjson", ".parquet", ".xlsx")


async def observe_data(path: str | Path) -> Observation:
    """
    Mira un conjunto de datos generado: filas, columnas y si tiene contenido.

    Segunda rama que faltaba en `observe()`. `ArtifactKind.DATA` estaba en el
    enum y el schema de `observe_artifact` ofrecía "datos", pero se despachaba
    a `observe_program`, es decir, se INTENTABA EJECUTAR el CSV. Mismo fallo
    que con vídeo y encontrado por el mismo test, que recorre el enum entero
    en vez de comprobar los tipos que se me ocurrieron.

    El modo de fallo que esto caza es el CSV con cabecera y cero filas: el
    fichero existe, se abre, tiene columnas con buen aspecto y no hay datos.
    """
    p = Path(path)
    if not p.exists():
        return Observation(False, ArtifactKind.DATA, "no existe",
                           problems=[f"{p} no existe"])

    problems: list[str] = []
    evidence: list[str] = [f"{p.stat().st_size:,} bytes"]
    filas = 0
    ext = p.suffix.lower()

    try:
        if ext in (".csv", ".tsv"):
            import csv as _csv
            with p.open("r", encoding="utf-8", errors="replace", newline="") as f:
                lector = _csv.reader(f, delimiter="\t" if ext == ".tsv" else ",")
                cabecera = next(lector, None)
                filas = sum(1 for _ in lector)
                if cabecera:
                    evidence.append(f"{len(cabecera)} columnas: "
                                    f"{', '.join(cabecera[:8])}")
                evidence.append(f"{filas:,} filas de datos")
        elif ext in (".jsonl", ".ndjson"):
            filas = sum(1 for ln in p.read_text(encoding="utf-8",
                                                errors="replace").splitlines()
                        if ln.strip())
            evidence.append(f"{filas:,} registros")
        elif ext == ".json":
            import json as _json
            datos = _json.loads(p.read_text(encoding="utf-8"))
            if isinstance(datos, list):
                filas = len(datos)
                evidence.append(f"lista de {filas:,} elementos")
            elif isinstance(datos, dict):
                filas = len(datos)
                evidence.append(f"objeto con {filas} claves: "
                                f"{', '.join(list(datos)[:8])}")
        else:
            # `.xlsx` y `.parquet` están en DATA_EXTS, así que `observe()` los
            # manda aquí — y aquí no se abrían. `filas = 1` por tener tamaño
            # hacía que el resumen AFIRMARA «1 registros» de un fichero que
            # nadie leyó: peor que el aviso, porque inventa el dato. Un
            # binario de 108 bytes de basura pasaba como conjunto de datos
            # válido.
            problems.append(
                f"formato {ext or 'sin extensión'} no se sabe abrir: el "
                f"contenido NO se ha mirado. Solo consta que el fichero pesa "
                f"{p.stat().st_size:,} bytes; el número de registros se "
                f"desconoce.")
            filas = -1
    except Exception as e:
        return Observation(False, ArtifactKind.DATA, "ilegible",
                           artifact_path=str(p),
                           problems=[f"no se pudo leer como {ext}: {e}"])

    if filas == 0:
        problems.append(
            "cero filas de datos: el fichero existe y tiene estructura, pero "
            "no contiene ningún registro. Es el fallo que más fácil pasa por "
            "bueno porque el fichero se abre sin errores.")

    resumen = (f"{p.name}: registros sin contar" if filas < 0
               else f"{p.name}: {filas:,} registros")
    return Observation(not problems, ArtifactKind.DATA, resumen,
                       evidence=evidence, artifact_path=str(p),
                       problems=problems)


async def observe_document(path: str | Path) -> Observation:
    """
    Cuenta páginas, palabras y detecta documentos vacíos.

    El fallo típico de generación de documentos es entregar un .docx con la
    plantilla y sin contenido, o un PDF de una sola página cuando se pidieron
    veinte.
    """
    p = Path(path)
    if not p.exists():
        return Observation(False, ArtifactKind.DOCUMENT, "no existe",
                           problems=[f"{p} no existe"])

    ext = p.suffix.lower()
    evidence, problems = [f"{p.stat().st_size:,} bytes"], []

    if ext == ".pdf":
        try:
            import pypdf
            r = pypdf.PdfReader(str(p))
            n = len(r.pages)
            text = "".join((pg.extract_text() or "") for pg in r.pages[:20])
            evidence.append(f"{n} páginas, {len(text.split())} palabras "
                            f"en las primeras {min(n, 20)}")
            if n == 0:
                problems.append("PDF sin páginas")
            elif len(text.strip()) < 40:
                problems.append("el PDF no tiene texto extraíble: ¿páginas en "
                                "blanco o solo imágenes?")
        except ImportError:
            # En `evidence` el aviso no entra en `ok`: un PDF de páginas en
            # blanco salía aprobado por no tener con qué abrirlo. Mismo fallo
            # que con Pillow, en otro formato.
            problems.append("pypdf no instalado: el PDF NO se ha inspeccionado. "
                            "No se sabe si tiene texto. `pip install pypdf`.")
        except Exception as e:
            problems.append(f"PDF ilegible: {e}")

    elif ext in (".docx", ".dotx"):
        try:
            import docx
            d = docx.Document(str(p))
            words = sum(len(par.text.split()) for par in d.paragraphs)
            evidence.append(f"{len(d.paragraphs)} párrafos, {words} palabras, "
                            f"{len(d.tables)} tablas")
            if words < 20:
                problems.append("el documento está prácticamente vacío")
        except ImportError:
            problems.append("python-docx no instalado: el documento NO se ha "
                            "inspeccionado. `pip install python-docx`.")
        except Exception as e:
            problems.append(f"docx ilegible: {e}")

    elif ext in (".md", ".txt", ".html"):
        text = p.read_text(encoding="utf-8", errors="replace")
        words = len(text.split())
        evidence.append(f"{words} palabras, {len(text.splitlines())} líneas")
        if words < 20:
            problems.append("prácticamente vacío")

    else:
        problems.append(f"tipo {ext or 'sin extensión'} no inspeccionable: el "
                        f"documento NO se ha mirado, así que no se puede decir "
                        f"que esté bien")

    if p.stat().st_size == 0:
        problems.append("fichero de 0 bytes")

    return Observation(not problems, ArtifactKind.DOCUMENT,
                       f"{p.name} ({ext or 'sin extensión'})",
                       evidence=evidence, artifact_path=str(p),
                       problems=problems)


# ----------------------------------------------------------------- común

async def _run(cmd: str, cwd: Path, timeout: int,
               extra_env: dict | None = None) -> tuple[int, str]:
    import os
    env = {**os.environ, **(extra_env or {})}
    cwd.mkdir(parents=True, exist_ok=True)
    try:
        proc = await asyncio.create_subprocess_shell(
            cmd, cwd=str(cwd), env=env,
            stdout=asyncio.subprocess.PIPE, stderr=asyncio.subprocess.STDOUT)
    except Exception as e:
        return 127, str(e)
    from ...core.cancel import tracked
    try:
        async with tracked(proc):
            out, _ = await asyncio.wait_for(proc.communicate(), timeout=timeout)
    except asyncio.TimeoutError:
        proc.kill()
        await proc.wait()
        return 124, f"timeout tras {timeout}s"
    return proc.returncode or 0, out.decode("utf-8", errors="replace")


async def observe(path: str | Path, kind: ArtifactKind | str | None = None,
                  **kw) -> Observation:
    """Despacha por tipo, deduciéndolo de la extensión si no se indica."""
    p = Path(path)
    if kind is None:
        ext = p.suffix.lower()
        if p.is_dir():
            kind = ArtifactKind.GAME
        elif ext in (".png", ".jpg", ".jpeg", ".webp", ".bmp"):
            kind = ArtifactKind.IMAGE
        elif ext in VIDEO_EXTS:
            kind = ArtifactKind.VIDEO
        elif ext in DATA_EXTS:
            kind = ArtifactKind.DATA
        elif ext in (".pdf", ".docx", ".md", ".txt", ".html", ".dotx"):
            kind = ArtifactKind.DOCUMENT
        else:
            kind = ArtifactKind.PROGRAM
    kind = ArtifactKind(kind) if isinstance(kind, str) else kind

    if kind is ArtifactKind.GAME:
        return await observe_game(p, **kw)
    if kind is ArtifactKind.IMAGE:
        return await observe_image(p)
    if kind is ArtifactKind.DOCUMENT:
        return await observe_document(p)
    if kind is ArtifactKind.VIDEO:
        # §5.5. Antes NO existía esta rama, aunque ArtifactKind.VIDEO estaba en
        # el enum y el schema de `observe_artifact` ofrecía "video" como valor
        # válido. Un .mp4 caía aquí abajo, en observe_program, y se intentaba
        # EJECUTAR como Python: el agente que pedía mirar el vídeo que acababa
        # de generar recibía "SyntaxError: source code cannot contain null
        # bytes". Una capacidad anunciada y no conectada es peor que una que
        # falta, porque nadie la busca.
        from .video import observe_video
        return await observe_video(p)
    if kind is ArtifactKind.DATA:
        return await observe_data(p)
    return await observe_program(p, **kw)


def available_backends() -> dict[str, bool]:
    """Qué se puede observar en esta máquina."""
    def has(mod: str) -> bool:
        import importlib.util
        return importlib.util.find_spec(mod) is not None

    return {
        "pygame": has("pygame"), "pillow": has("PIL"),
        "pypdf": has("pypdf"), "python-docx": has("docx"),
        "ffmpeg": bool(shutil.which("ffmpeg")),
        # ffprobe va aparte de ffmpeg a propósito: algunas distribuciones lo
        # empaquetan separado, y sin él se pueden GENERAR vídeos pero no
        # inspeccionarlos — es decir, se rompe el bucle de observación
        # justamente donde no se nota.
        "ffprobe": bool(shutil.which("ffprobe")),
        "comfyui_local": _comfy_reachable(),
    }


def backends_report() -> str:
    """Qué se puede hacer aquí, y qué falta para lo demás."""
    b = available_backends()
    lines = [f"  {'sí' if v else 'no':<4s} {k}" for k, v in b.items()]
    notas = []
    if not b["pygame"]:
        notas.append("sin pygame no se pueden observar juegos, solo leer su código")
    if not b["comfyui_local"]:
        notas.append("sin ComfyUI en 127.0.0.1:8188 la composición de manga "
                     "funciona pero las viñetas salen como marcadores de "
                     "posición, no como dibujos")
    if not b["ffmpeg"]:
        notas.append("sin ffmpeg no hay vídeo programático")
    elif not b["ffprobe"]:
        notas.append("hay ffmpeg pero no ffprobe: se pueden generar vídeos y "
                     "NO inspeccionarlos, así que no se detectaría uno negro "
                     "o congelado")
    return "\n".join(lines) + ("\n\n" + "\n".join(f"- {n}" for n in notas)
                               if notas else "")


def _comfy_reachable(host: str = "http://127.0.0.1:8188") -> bool:
    """ComfyUI para imagen/manga (§5.4). Local, gratis, sin claves."""
    try:
        import urllib.request
        with urllib.request.urlopen(f"{host}/system_stats", timeout=1):
            return True
    except Exception:
        return False
